# Standard library imports
import csv
import json
from math import e
import os
import re
import traceback
import logging
import logging.config
import yaml
import docx
from typing import List, Tuple
from xml.dom import minidom
from docx import Document, document, table,opc,oxml
from docx.table import _Cell, Table
from docx.shared import Pt,Cm
import xmlschema
import unicodedata
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Third-party imports
from bs4 import BeautifulSoup

# Constants and Configuration
nexo_tool_config:dict[str,str]={}
ISO20022_MODEL = 'C:/Users/romai/OneDrive/Documents/NEXO/api/20240729_ISO20022_2013_eRepository.iso20022'
REPO_DESC_FILE = "C:/Users/romai/OneDrive/Documents/NEXO/api/" + \
    os.path.basename(ISO20022_MODEL) + ".description.json"
PREVIOUS_VERSION_NUMBER = "02"
ROW_LEVEL_INDENT = "  "

LOG_CONFIGURATION_FILE = "nexo_table_generator_logging.yaml"

OUTPUT_DIR = "C:/Users/romai/OneDrive/Documents/NEXO/api/casp_v5/NEW_HTML/ENDPOINTS_WITHOUT_ENDPOINTS/"
SCHEMA_DIR = "C:/Users/romai/OneDrive/Documents/NEXO/api/casp_v5/"
PREVIOUS_VERSION_DIR = "C:/Users/romai/OneDrive/Documents/NEXO/api/casp_v5/HTML/"
V5_VERSION_DIR = "C:/Users/romai/OneDrive/Documents/NEXO/api/casp_v2/NEW_HTML/"
# Set this to False if you don't want to generate endpoint tables
GENERATE_ENDPOINTS = True
# Set this to False if you don't want to see stop at sub endpoints in endpoints
ENDPOINTS_IN_ENDPOINT=True
# Set this to False if you don't want to compare message to a previous version
COMPARE_TO_PREVIOUS_VERSION = True
# Set this to False if you don't want to compare endpoints to a previous version
ENDPOINT_COMPARE_TO_PREVIOUS_VERSION = True
SPECIAL_COMPARE_V5 = True


#OUTPUT_DIR = "C:/Users/romai/OneDrive/Documents/NEXO/api/casp_v2/NEW_HTML/"
#SCHEMA_DIR = "C:/Users/romai/OneDrive/Documents/NEXO/api/casp_v2/"
#PREVIOUS_VERSION_DIR = "C:/Users/romai/OneDrive/Documents/NEXO/api/casp_v5/HTML/"
#V5_VERSION_DIR = "C:/Users/romai/OneDrive/Documents/NEXO/api/casp_v2/NEW_HTML/"
###Set this to False if you don't want to generate endpoint tables
#GENERATE_ENDPOINTS = True
##Set this to False if you don't want to see stop at sub endpoints in endpoints
#ENDPOINTS_IN_ENDPOINT=False
##Set this to False if you don't want to compare message to a previous version
#COMPARE_TO_PREVIOUS_VERSION = False
#ENDPOINT_COMPARE_TO_PREVIOUS_VERSION = False
#SPECIAL_COMPARE_V5 = False

# Status Constants
STATUS_NOCHANGE = 0
STATUS_UPDATED = 1
STATUS_NEW = 2
STATUS_DELETED = 4

# Global variables

logger: logging.Logger
endpoints: dict[str, str] = {}
business_area: str = ""
iso_repository_description: dict[str, dict[str, str]] = {}
bs_previous_html: BeautifulSoup
element_tags: dict[str, str] = {}
message_tags: dict[str, dict[str,str]] = {}
endpoints_elements: dict[str, xmlschema.validators.elements.XsdElement] = {}

# XML namespaces
namespaces = {
    "xmi": "http://www.omg.org/XMI",
    "xs": "http://www.w3.org/2001/XMLSchema"
}

# HTML Style Template
HTML_STYLE = """<style type="text/css">
    p {
        color: #000000;
        line-height: 95%;
        text-align: left;
        orphans: 2;
        widows: 2;
        margin-bottom: 0.25cm;
        direction: ltr;
        background: transparent
    }
    p.western {
        font-family: "Arial", serif;
        font-size: 10pt;
        so-language: en-GB
    }
    a:visited { color: #954f72; text-decoration: underline }
    a:link { color: #0563c1; text-decoration: underline }
    table.fixed {
        table-layout: fixed;
        width: 840px;
        border-collapse: collapse;
        border: 1px solid #a6a6a6;
    }
    tr.updated {
        font-color: #0000ff;
    }
    tr.new {
        font-color: #00cc00;
    }
    tr.deleted {
        font-color: #ff0000;
        text-decoration: line-through;
    }

    table.fixed td { word-wrap: break-word; }
    table.fixed td:nth-of-type(1) { width: 32px; }
    table.fixed td:nth-of-type(2) { width: 190px; }
    table.fixed td:nth-of-type(3) { width: 60px; }
    table.fixed td:nth-of-type(4) { width: 60px; }
    table.fixed td:nth-of-type(5) { width: 60px; }
    table.fixed td:nth-of-type(6) { width: 438px; }
</style>"""

HTML_BODY = '''
    <body>
    '''
namespaces = {"xmi": "http://www.omg.org/XMI",
              "xs": "http://www.w3.org/2001/XMLSchema"}


def load_endpoints() -> dict[str, str]:
    """
    Load endpoints from CSV file.
    Returns a dictionary mapping element types to protocols.
    """
    mdr_endpoints: dict[str, str] = {}
    for filename in os.listdir(SCHEMA_DIR):
        if filename.endswith(".csv") and (filename.upper().startswith("NEW_EXPANDABLES")):
            logger.info(f"Processing {filename} for endpoints")
            with open(os.path.join(SCHEMA_DIR, filename), 'r', encoding='UTF-8') as file:
                csv_reader = csv.reader(file, delimiter=";")
                for elt_type, protocol in csv_reader:
                    mdr_endpoints[elt_type] = protocol
    return mdr_endpoints


def is_endpoint(elt_type: str = "") -> bool:
    """
    Check if the given element type is an endpoint for the current business area.
    """
    return elt_type in endpoints and business_area in endpoints[elt_type]


def collect_all_endpoints_elements(doc, endpoints_dict:dict[str, xmlschema.validators.elements.XsdElement]) -> dict[str, xmlschema.validators.elements.XsdElement]:
    """
    iterates over a XMLSchema and collect all elements that are endpoints.
    Args:
        doc : a XMLSchema Element document root
        endpoints_dict : a dictionary mapping element types to XsdElement objects for endpoints.

    Returns:
         updated dictionary mapping element types to XsdElement objects for endpoints.
    """
    for element in doc.iterchildren():
        if element.type is not None and is_endpoint(element.type.local_name): # type: ignore
            endpoints_dict[element.type.local_name] = element 
        endpoints_dict = collect_all_endpoints_elements(
            element, endpoints_dict)
    return endpoints_dict


def generate_word_message_table_xsd_endpoint(message_name: str, output_dir: str, element_type: str, endpoint_elt) -> None:
    """
    Generate HTML table for a endpoint.
    Depending on compare flags, it will try to load HTML files from another version of the schema and compare the elements
    to highlight changes between versions.

    Args:
        message_name: Name of the message to process
        output_dir: Directory to write output HTML files
        element_type: type of element to process
        endpoint_elt : XsdElement of the endpoint (typically collected in the full message table geenration)

    """

    table_title = message_name + ("." + element_type)

    previous_html: BeautifulSoup = BeautifulSoup()
    v5_html: BeautifulSoup = BeautifulSoup()
    if (COMPARE_TO_PREVIOUS_VERSION):
        try:
            previous_html = load_previous_version_html(
                message_name, element_type)
        except Exception as preve:
            logger.error(f"Error finding previous version HTML for {
                element_type}: {preve}")

    if (SPECIAL_COMPARE_V5):
        try:
            v5_html = load_v5_version_html(
                message_name, element_type)
        except Exception as preve:
            logger.error(f"Error loading v5 version HTML for {
                element_type}: {preve}")

    # Generate table content
    document,table = generate_word_table_header(element_type)
    table = process_word_endpoint_elements(document, table,
        endpoint_elt, 1, previous_html, v5_html)



    # Write output file
    output_filename = f"{
        element_type}.docx"
    output_file = os.path.join(output_dir, output_filename)
    set_col_widths(table)
    document.save(output_file)
    logger.info("-" * 16 + f"Endpoint file written: {output_file}")


def generate_word_message_table_from_schema(input_dir: str, xs: xmlschema.XMLSchema, output_dir: str) -> None:
    """
    Generate word table for a specific message or sub-element.
    Args:
        input_dir: Directory containing input schema files
        xs: XML Schema object
        output_dir: Directory to write output word files
    """

    # Process schema
    logger.info(f"Processing {xs.name}")
    previous_html: BeautifulSoup = BeautifulSoup()
    message_title = "_MESSAGE_TITLE_"

    # Check if 'Document' element exists
    if 'Document' in xs.elements:
        business_area = xs.name[0:4]  # type: ignore
        doc = xs.elements['Document']
        message_title = xs.name if xs.name is not None else message_title
        v5_html: BeautifulSoup = BeautifulSoup()

        try:
            for elt in doc.iterchildren():
                message_title = element_tags[elt.local_name] if (
                    elt.local_name in element_tags) else message_title
                break
        except Exception as ex:
            logger.info(f"Error while getting message title for {
                        elt.local_name}: {ex}")

        if (COMPARE_TO_PREVIOUS_VERSION):
            try:
                # remove version of the message for comparisons (StatusReportV07 => StatusReport)
                previous_html = load_previous_version_html(
                    re.sub(r'V[0-9]+$', '', message_title), "")  # type: ignore
            except Exception as preve:
                logger.error(f"Error loading previous version HTML for {
                    message_title}: {preve}")

        if (SPECIAL_COMPARE_V5):
            try:
                v5_html = load_v5_version_html(
                    re.sub(r'V[0-9]+$', '', message_title), "")
            except Exception as preve:
                logger.info(f"Error loading v5 version HTML for {
                    message_title}: {preve}")
        # Generate table content
        document,table = generate_word_table_header(
            re.sub(r'V[0-9]+$', '', message_title))  # type: ignore

        table = process_word_schema_elements(document,table,message_title,
                                                  doc, level=1, previous_version=previous_html, v5_version=v5_html)


      
        # Write output file
        message_title = re.sub(r'V[0-9]+$', '', message_title)  # type: ignore

        output_filename = f"{message_title}.docx"
        output_file = os.path.join(output_dir, output_filename)
        set_col_widths(table)
        document.save(output_file)

    else:
        logger.error(f"No 'Document' element found in schema {xs.name}")



def find_schema_file(input_dir: str, message_name: str) -> str:
    """
    Find the schema file for a given message name.

    Args:
        input_dir: Directory to search in
        message_name: Name of the message to find ex: casp.002.001.06

    Returns:
        Full path to schema file or empty string if not found
    """
    for file in os.listdir(input_dir):
        if file.startswith(message_name):
            return os.path.join(input_dir, file)
    return ""



def set_col_widths(table):
    widths = (Cm(1), Cm(6), Cm(1.5), Cm(1.5), Cm(1.5), Cm(6.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def generate_word_table_header(head_component: str) :
    """
    Generate the HTML table header.

    Args:
        head_component: Document element from schema (Ex: SaleToPOIServiceResponseV06)

    Returns:
        HTML string containing table header
    """
    document = Document()
    document.add_heading(head_component, 0)
    table = document.add_table(rows=1, cols=6)
    table.autofit = False 
    table.allow_autofit = False
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Lvl'
    hdr_cells[0].width = Cm(1)
    hdr_cells[1].text = head_component
    hdr_cells[1].width = Cm(6)
    hdr_cells[2].text = 'Mult'
    hdr_cells[2].width = Cm(1.5)
    hdr_cells[3].text = 'Constraint'
    hdr_cells[3].width = Cm(1.5) 
    hdr_cells[4].text = 'Rule'
    hdr_cells[4].width = Cm(1.5)
    hdr_cells[5].text = 'Usage'
    hdr_cells[4].width = Cm(6.5)
    for cell in hdr_cells:
        set_cell_border(
            cell,
            top={"sz": 12, "color": "#000000", "val": "single"},
            bottom={"sz": 12, "color": "#000000", "val": "single"},
            start={"sz": 12, "color": "#000000", "val": "single"},
            end={"sz": 12, "color": "#000000", "val": "single"}
        )

    return document,table



def load_previous_version_html(message_name: str, element_type: str) -> BeautifulSoup:
    """
    Load and parse the HTML file from the previous version.
    Args:
        message_name: Name of the message to load
    Returns:
        BeautifulSoup object containing the parsed HTML
    """

    if (element_type == ''):
        with open(file=os.path.join(PREVIOUS_VERSION_DIR, message_name+".html"), mode='r', encoding="UTF-8") as file:
            soup = BeautifulSoup(markup=file, features='lxml')
    else:
        # Remove digits from the element type contentInformationType33 => contentInformationType
        # TODO : improve using previous version attribute in repository
        element_type_without_digits = re.sub(r'\d+$', '', element_type)
        if (element_type + ".html") in os.listdir(PREVIOUS_VERSION_DIR):
            logger.debug(f"Found same version in {os.path.join(PREVIOUS_VERSION_DIR, element_type + ".html")} for {element_type}")
            with open(file=os.path.join(PREVIOUS_VERSION_DIR, element_type + ".html"), mode='r', encoding="UTF-8") as file:
                soup = BeautifulSoup(markup=file, features='lxml')
        else:
            # try to find anotther version of the element\_type
            for filename in os.listdir(PREVIOUS_VERSION_DIR):
                # Check if the filename (without extension) matches the element type without digits
                if filename.lower().startswith(element_type_without_digits.lower()) and filename.lower().endswith('.html'):
                    logger.debug("Found another version for " +
                                 element_type + " ; Using file " + filename)
                    with open(file=os.path.join(PREVIOUS_VERSION_DIR, filename), mode='r', encoding="UTF-8") as file:
                        soup = BeautifulSoup(markup=file, features='lxml')
                    break
            logger.info("No previous version found for " + element_type)    

    return soup


def load_v5_version_html(message_name: str, element_type: str) -> BeautifulSoup:
    """
    Load and parse the HTML file from the v5 specific version.
    Args:
        message_name: Name of the message to load
    Returns:
        BeautifulSoup object containing the parsed HTML
    """

    if (element_type == ''):
        with open(file=os.path.join(V5_VERSION_DIR, message_name+".html"), mode='r', encoding="UTF-8") as file:
            soup = BeautifulSoup(markup=file, features='lxml')
    else:
        # Remove digits from the element type contentInformationType33 => contentInformationType
        # TODO : improve using previous version attribute in repository
        element_type_without_digits = re.sub(r'\d+$', '', element_type)
        if (element_type + ".html") in os.listdir(V5_VERSION_DIR):
            logger.debug("v5: Found same version in "+ V5_VERSION_DIR +" for " + element_type)
            with open(file=os.path.join(V5_VERSION_DIR, element_type + ".html"), mode='r', encoding="UTF-8") as file:
                soup = BeautifulSoup(markup=file, features='lxml')
        else:
            # try to find anotther version of the element\_type
            for filename in os.listdir(V5_VERSION_DIR):
                # Check if the filename (without extension) matches the element type without digits
                if filename.lower().startswith(element_type_without_digits.lower()) and filename.lower().endswith('.html'):
                    logger.debug("V5 : Found another version for " +
                                 element_type + " ; Using file " + V5_VERSION_DIR +filename)
                    with open(file=os.path.join(V5_VERSION_DIR, filename), mode='r', encoding="UTF-8") as file:
                        soup = BeautifulSoup(markup=file, features='lxml')
                    break
    return soup


def parse_repository(iso2002_repository_file: str) -> tuple[dict[str, str], dict[str, dict[str, str]]]:
    """
    use swift repository to associate message component names and xmlTags
    needed because XML schema 2023 does not contain Long name
    (not needed in XML Schema 2024 with annotations)
    Args:
        message_name: ISO20022 repository file

    Returns:
        Dictionary with xml_tags and name for elements, and message definitions
        Dictionary associating message name with message building blocks tag/names {'SaleToPOIServiceRequestV06',{'SctyTrlr','SecurityTrailer'}}

    """
    logger.info("Parsing repository for messageElement and messageDefinition")
    root = minidom.parse(file=iso2002_repository_file)

    elements = root.getElementsByTagName('messageElement')
    msg_defs = root.getElementsByTagName('messageDefinition')
    message_tags: dict[str, dict[str, str]] = {}

    for elt in elements:
        # logger.debug("EltName: "+elt.attributes['name'].value)
        if 'xmlTag' in elt.attributes and 'name' in elt.attributes:
            # logger.debug(elt.attributes['name'].value+"="+elt.attributes['xmlTag'].value)
            element_tags[elt.attributes['xmlTag'].value] = elt.attributes['name'].value
    for msg in msg_defs:
        # logger.debug("MessageName: "+elt.attributes['name'].value)
        if 'name' in msg.attributes and 'xmlTag' in msg.attributes:
            # logger.debug(msg.attributes['name'].value+"="+msg.attributes['xmlTag'].value)
            element_tags[msg.attributes['xmlTag'].value] = msg.attributes['name'].value
            message_blocks = msg.getElementsByTagName('messageBuildingBlock')
            mbbs:dict[str,str] = {}
            #for each message get message Building blocks tags and names
            for mbb in message_blocks:
                if 'name' in mbb.attributes and 'xmlTag' in mbb.attributes:
                    mbbs[mbb.attributes['xmlTag'].value]= mbb.attributes['name'].value
            message_tags[msg.attributes['name'].value]=mbbs

    # logger.debug(sizeof(xmlTags.values).)
    logger.info(
        "================================================================")
    # Read in the file where xmlTags must be replaced with the full name of the message component
    return element_tags, message_tags


def previous_version_lookup(previous_version: BeautifulSoup, elt_name: str, elt_type: str, elt_mult: str) -> tuple[int, str, str, str]:
    """
    Compare current version with previous version to determine if changes or new.
    Args:
        elt_name: Element name to look up
        elt_type: Element type
        elt_mult: Element multiplicity
    Returns:
        Tuple containing (status, previous_usage, rule, constraint)
    """
    status: int = STATUS_NOCHANGE
    found: bool = False
    prev_usage: str = ''  # '<b><font color="#d35400">TO VERIFY!!</font></b>'
    rule: str = ""
    constraint: str = ""

    if previous_version is not None:
        for row in previous_version.findAll(name="tr"):
            cells = row.findAll("td")
            level: str = cells[0].text.strip()
            name: str = cells[1].text.strip()
            mult: str = cells[2].text.strip()
            rule: str = cells[3].text.strip()
            constraint: str = cells[4].text.strip()
            usage = ""
            #usage may have html tags inside of it so we need to agregates all contents
            usage = cells[5].decode_contents().replace("\n", "").replace("\t", "")

            usage: str = re.sub(
                r'<br/>$', ' ', usage)
            
            #not needed anymore
            #usage=re.sub(r"&nbsp([^;])", '\\1', usage)
            if (elt_type in usage) and (elt_name in name):
                # logger.info(f"Found {elt_name} - Usage {usage}")
                prev_usage = usage
                # Check if multiplicity has changed
                if elt_mult != mult:
                    logger.debug(f"Found {elt_name} - Usage {usage}")
                    status = STATUS_UPDATED
                found = True
                break
            # try to identify update of Elements (as GenericIdentificationXXX) and CodeSet (like CardPaymentServiceTypeXXCode )
            if ((re.sub('[0-9]+', '', elt_type) in usage) or (re.sub('[0-9]+Code', '', elt_type) in usage)) and (elt_name in name):
                logger.debug(f"!!!!!!!!!!!!!!!!!!Found UPDATE {elt_name} {
                             elt_type}- Previous Usage {usage}")
                # prev_usage = usage + f' TO VERIFY <br/> UPDATE TO {elt_type}  '
                status = STATUS_UPDATED
                found = True
                break
        if not (found):
            logger.debug(f"!!!!!!!!!!!!!!!!!!NEW {elt_name} {elt_type}")
            found = True
            status = STATUS_NEW

        # clean previous ref to endpoint
        prev_usage: str = re.sub(r'See(.*)$', '', prev_usage)

    return status, prev_usage, rule, constraint


def V5_special_lookups(v5_version: BeautifulSoup, elt_name: str, elt_type: str, elt_mult: str) -> int:
    """
    special compare with v5 to update status only

    Returns:
        status
    """
    status: int = STATUS_NOCHANGE
    found: bool = False

    if v5_version is not None:
        for row in v5_version.findAll(name="tr"):
            cells = row.findAll("td")
            level: str = cells[0].text.strip()
            name: str = cells[1].text.strip()
            mult: str = cells[2].text.strip()
            usage: str = cells[5].text.strip()

            if (elt_type in usage) and (elt_name in name):
                # logger.info(f"Found {elt_name} - Usage {usage}")
                # Check if multiplicity has changed
                if elt_mult != mult:
                    logger.debug(f"!!!!!!!!!!!!!!!!!!F V5 UPDATE Found {elt_name} {
                                 elt_mult} changed from {mult}- Usage {usage}")
                    status = STATUS_UPDATED
                found = True
                break
            # try to identify update of Elements (as GenericIdentificationXXX) and CodeSet (like CardPaymentServiceTypeXXCode )
            if ((re.sub('[0-9]+', '', elt_type) in usage) or (re.sub('[0-9]+Code', '', elt_type) in usage)) and (elt_name in name):
                logger.debug(f"!!!!!!!!!!!!!!!!!!Found UPDATE from V5 for {
                             elt_name} {elt_type}- Previous {usage}")
                # not needed anymore
                # prev_usage = usage + f' TO VERIFY <br/> UPDATE TO {elt_type}  '
                status = STATUS_UPDATED
                found = True
                break
        if not (found):
            logger.debug(f"!!!!!!!!!!!!!!!!!!NEW {elt_name} {
                         elt_type} NOT FOUND IN V5")
            found = True
            status = STATUS_NEW



    return status




def process_word_schema_elements(document, table, message_name: str, xsd_element, level: int, previous_version: BeautifulSoup, v5_version: BeautifulSoup) -> table:
    """
    Recursively process and format element children into word table rows.

    Args:
        xsd_element: Current element being processed
        level: Current nesting level
        previous_version: BeautifulSoup object containing the previous version HTML
        v5_version: BeautifulSoup object containing the v5 version HTML
    Returns:
        HTML string containing formatted table rows
    """
    html_table_rows = ""

    # Extract element metadata
    element_info = get_element_info(message_name, xsd_element)

    # Handle special case where name ends with a number or V+Number such as ReversalResponseV03 or ServiceResponse7
    element_info['name'] = re.sub(r'V[0-9]+$', '', element_info['name'])
    element_info['name'] = re.sub(r'[0-9]+$', '', element_info['name'])

    # Process element multiplicity
    mult_info = get_multiplicity_info(xsd_element)

    elt_status = STATUS_NOCHANGE
    elt_rule = ""
    elt_constraint = ""
    elt_name: str = element_info['name']
    elt_type = element_info['type']
    elt_mult = mult_info['multiplicity']
    elt_def = element_info['definition']
    elt_tag = element_info['tag']

    elt_usage = ""

    # Handle special cases and generate row HTML
    if level > 2:
        if (COMPARE_TO_PREVIOUS_VERSION):
            elt_status, prev_usage, elt_rule, elt_constraint = previous_version_lookup(previous_version,
                                                                                       elt_name, elt_type, elt_mult)
            elt_usage = elt_usage + \
                clean_previous_usage(prev_usage, elt_type)

        if is_endpoint(elt_type) and elt_type:
            linebreak: str = "<br/>" if elt_usage else ""
            elt_usage = f'{elt_usage}{linebreak}See MDR for sub elements and <a href="#{
                elt_type}">{elt_type}</a>'

        row_color = get_status_color(elt_status)

        if (SPECIAL_COMPARE_V5):
            elt_status = V5_special_lookups(
                v5_version, elt_name, elt_type, elt_mult)
            row_color = get_status_color(elt_status)

        table = generate_word_table_row(document, table,
            level=level-2,
            element_info=element_info,
            mult_info=mult_info,
            prev_info={
                'usage': elt_usage,
                'rule': elt_rule,
                'constraint': elt_constraint
            },
            color=row_color,
            indent_level=level
        )
    
    # Recursively process children if not an endpoint
    if element_info['type']:
        for xsd_component in xsd_element.iterchildren():
            try:
                if (xsd_component.local_name is not None) and not (is_endpoint(element_info['type'])):
                    table = process_word_schema_elements(document,table,message_name,
                                                               xsd_component, level + 1, previous_version, v5_version)
                else:
                    logger.debug(msg=f"endpoint reached: {
                                 element_info['type']}")
            except TypeError as te:
                logger.error(f"TypeError: {te}")

    return table


def process_word_endpoint_elements(document, table, xsd_element,level, previous_version, v5_version: BeautifulSoup) -> table:
    """
    Recursively process and format element children into HTML table rows.

    Args:
        elt: Current element being processed
        level: Current nesting level
        previous_version: BeautifulSoup object containing the previous version HTML
        v5_version: BeautifulSoup object containing the v5 version HTML
    Returns:
        HTML string containing formatted table rows
    """
   
    # Extract element metadata
    #we are in a endpoint so metadat is not related to a message
    element_info = get_element_info("",xsd_element)

    # Handle special case where name ends with a number or V+Number such as ReversalResponseV03 or ServiceResponse7
    element_info['name'] = re.sub(r'V[0-9]+$', '', element_info['name'])
    element_info['name'] = re.sub(r'[0-9]+$', '', element_info['name'])

    # Process element multiplicity
    mult_info = get_multiplicity_info(xsd_element)

    elt_status = STATUS_NOCHANGE
    elt_rule = ""
    elt_constraint = ""
    elt_name: str = element_info['name']
    elt_type = element_info['type']
    elt_mult = mult_info['multiplicity']
    elt_def = element_info['definition']
    elt_tag = element_info['tag']
    
    #skip the 1st element which is the endpoint itself

    if (level>1):
        elt_usage = ""
        if (COMPARE_TO_PREVIOUS_VERSION):
            elt_status, prev_usage, elt_rule, elt_constraint = previous_version_lookup(previous_version,
                                                                                    elt_name, elt_type, elt_mult)
            elt_usage = elt_usage + \
                clean_previous_usage(prev_usage, elt_type)
        
        if is_endpoint(elt_type) and elt_type and ENDPOINTS_IN_ENDPOINT:
            elt_usage = f'{elt_usage}<br>See MDR for sub elements and <a href="#{
                elt_type}">{elt_type}</a>'
        
        row_color = get_status_color(elt_status)
        
        if (SPECIAL_COMPARE_V5):
            elt_status = V5_special_lookups(
                v5_version, elt_name, elt_type, elt_mult)
            row_color = get_status_color(elt_status)
    
        table = generate_word_table_row(document,table,
            level=level-1,
            element_info=element_info,
            mult_info=mult_info,
            prev_info={
                'usage': elt_usage,
                'rule': elt_rule,
                'constraint': elt_constraint
            },
            color=row_color,
            indent_level=level
        )



    # Recursively process children
    for xsd_component in xsd_element.iterchildren():
        try:
            if (xsd_component.type.local_name is not None) and not ((is_endpoint(element_info['type']) and ENDPOINTS_IN_ENDPOINT) and (level > 1)):
                table = process_word_endpoint_elements(document,table,
                    xsd_component, level + 1, previous_version,v5_version)
            else:
                logger.debug(msg=f"endpoint reached: {element_info['type']}")
        except TypeError as te:
            logger.error(f"TypeError: {te}")

    return table


def get_element_info_XML2024(elt) -> dict:
    """
    Extract basic information from an element.
    NEDS NEW XML SCHEMA 2024


    Args:
        elt: Element to process

    Returns:
        Dictionary containing element information
    """
    info = {
        'name': ' - ',
        'definition': ' - ',
        'tag': ' - ',
        'type': ' - '
    }

    if len(elt.annotations) > 0:
        info['name'] = elt.annotations[0].documentation[1].text
        info['definition'] = elt.annotations[0].documentation[0].text
        info['tag'] = elt.local_name

    if 'type' in elt.schema_elem.attrib:
        info['type'] = elt.elem.attrib['type']

    if info['type'].endswith("Code"):
        info['definition'] = (info['definition'] +
                              iso_repository_description[info['type']]['def'])

    return info


def get_element_info(message_name,elt) -> dict[str, str]:
    """
    Extract basic information from an element.

    Args:
        message_name : a message name like StatusReportV07
        elt: XML schema Element to process

    Returns:
        Dictionary containing element information
    """
    info = {
        'name': ' - ',
        'definition': ' - ',
        'tag': ' - ',
        'type': ' - '
    }
    try:
        if getattr(elt,'type'):
            info['type'] = getattr(elt.type, 'local_name')
            #Context: endpoint
            if (message_name == ""):
                info['name'] = element_tags[getattr(
                    elt, 'local_name')] if getattr(elt, 'local_name') in element_tags else info['type']
            else:
                #Context : message
                info['name'] = message_tags[message_name][getattr(
                    elt, 'local_name')] if getattr(elt, 'local_name') in message_tags[message_name] else element_tags[getattr(
                        elt, 'local_name')] if getattr(elt, 'local_name') in element_tags else info['type']

            info['definition'] = iso_repository_description[info['type']
                                                            ]['def'] if info['type'] in iso_repository_description else ' - '
            info['tag'] = getattr(elt, 'local_name')
    except AttributeError as e:
        logger.error(f"AttributeError: {e} - {elt} - {message_name}")

    return info


def get_multiplicity_info(elt) -> dict:
    """
    Extract multiplicity information from an element and format it [min..max]

    Args:
        elt: XML schema Element to process

    Returns:
        Dictionary containing multiplicity information
    """
    max_occurs = '*' if elt.effective_max_occurs is None else elt.effective_max_occurs
    return {
        'min': elt.effective_min_occurs,
        'max': max_occurs,
        'multiplicity': f'[{elt.effective_min_occurs}..{max_occurs}]'
    }


def get_status_color(status: int) -> str:
    """
    Get the color code for a given status.

    Args:
        status: Status code (NOCHANGE, UPDATED, NEW, or DELETED)

    Returns:
        Hex color code string
    """
    status_colors = {
        STATUS_NOCHANGE: "#000000",
        STATUS_UPDATED: "#28b463",
        STATUS_DELETED: "#5d6d7e",
        STATUS_NEW: "#2874a6"
    }
    return status_colors.get(status, "#000000")


def clean_previous_usage(usage: str, elt_type: str) -> str:
    """
    Clean up the previous usage text by removing unnecessary elements.

    Args:
        usage: Original usage text
        elt_type: Element type

    Returns:
        Cleaned usage text
    """
    
    # Remove <tag>::type pattern
    usage = re.sub(r"<i>&lt;\w+&gt;::\w+<\/i>", '', usage)
    # Remove useless sentence when endpoint
    # not needed anymore
    # usage = re.sub(r"See MDR(.*)trouvée.", '', usage)

    # Remove Initial line break if any
    usage = re.sub(r"<br/>$", ' ', usage)
    # Remove ending line break if any
    # # not needed anymore
    # usage = re.sub(r"^^<br/>", ' ', usage)
    #sometimes we find non breaking spaces misssing the semi comma ;
    #not needed anymore
    #usage = re.sub(r"&nbsp([^;])", ' \\1', usage)
    return usage


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)


    return hyperlink

def html_to_word(paragraph, html:str):
    """
    Convert HTML content to Word format and append it to a paragraph.
    """
    htmls = html.split('<')
    htmls = [html[0]] + ['<'+l for l in htmls[1:]]
    tags = []
    for run in htmls:
        tag_change = re.match('(?:<)(.*?)(?:>)', run)
        if tag_change != None:
            tag_strip = tag_change.group(0)
            tag_change = tag_change.group(1)
            if tag_change.startswith('/'):
                if tag_change.startswith('/a'):
                    tag_change = next(tag for tag in tags if tag.startswith('a '))
                tag_change = tag_change.strip('/')
                tags.remove(tag_change)
            else:
                tags.append(tag_change)
        else:
            tag_strip = ''
        hyperlink = [tag for tag in tags if tag.startswith('a ')]
        if run.startswith('<'):
            run = run.replace(tag_strip, '')
            if hyperlink:
                hyperlink = hyperlink[0]
                hyperlink = re.match(
                    '.*?(?:href=")(.*?)(?:").*?', hyperlink).group(1)
                add_hyperlink(paragraph, run, hyperlink)
            else:
                runner = paragraph.add_run(run)
                if 'b' in tags:
                    runner.bold = True
                if 'u' in tags:
                    runner.underline = True
                if 'i' in tags:
                    runner.italic = True
                if 'H1' in tags:
                    runner.font.size = Pt(24)
        else:
            paragraph.add_run(run)

    return paragraph


def generate_word_table_row(document,table,level: int, element_info: dict, mult_info: dict,
                       prev_info: dict, color: str, indent_level: int) -> table:
    indent = multStr(ROW_LEVEL_INDENT, indent_level)
    if (prev_info['usage'] != ""):
        prev_info['usage'] = prev_info['usage']+"<br/>"

    # add ISO Definition for codeset if usage does not precise a single value allowed
    if re.match(r'(.*[0-9]+)Code', element_info['type']) and not ('value must' in (prev_info['usage']).lower()):
        prev_info['usage'] = prev_info['usage']+"<br/>" + \
            iso_repository_description[element_info['type']
                                       ]['def']
    row_cells = table.add_row().cells
    row_cells[0].text=str(level)
    set_cell_border(
        row_cells[0],
        top={"sz": 12, "color": "#000000", "val": "single"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 12, "color": "#000000", "val": "single"},
        end={"sz": 12, "color": "#000000", "val": "single"}
    )
    row_cells[1].text=indent+element_info['name']
    set_cell_border(
        row_cells[1],
        top={"sz": 12, "color": "#000000", "val": "single"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 12, "color": "#000000", "val": "single"},
        end={"sz": 12, "color": "#000000", "val": "single"}
    )
    row_cells[2].text=mult_info['multiplicity']
    set_cell_border(
        row_cells[2],
        top={"sz": 12, "color": "#000000", "val": "single"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 12, "color": "#000000", "val": "single"},
        end={"sz": 12, "color": "#000000", "val": "single"}
    )
    row_cells[3].text=prev_info['rule']
    set_cell_border(
        row_cells[3],
        top={"sz": 12, "color": "#000000", "val": "single"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 12, "color": "#000000", "val": "single"},
        end={"sz": 12, "color": "#000000", "val": "single"}
    )
    row_cells[4].text=prev_info['constraint']
    set_cell_border(
        row_cells[4],
        top={"sz": 12, "color": "#000000", "val": "single"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 12, "color": "#000000", "val": "single"},
        end={"sz": 12, "color": "#000000", "val": "single"}
    )
    #row_cells[5].text=prev_info['usage']+"<i>&lt;"+element_info['tag']+"&gt;::"+element_info['type']+"</i>"
    p = row_cells[5].add_paragraph()
    p = html_to_word(p, prev_info['usage']+"<i><"+element_info['tag']+">::"+element_info['type']+"</i>")
    p.add_run(prev_info['usage'])
    rtag= p.add_run("<"+element_info['tag']+">::"+element_info['type'])
    rtag.italic = True
    set_cell_border(
        row_cells[5],
        top={"sz": 12, "color": "#000000", "val": "single"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 12, "color": "#000000", "val": "single"},
        end={"sz": 12, "color": "#000000", "val": "single"}
    )

    return table


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def multStr(char: str, n: int) -> str:
    """
    Create a string by repeating a character n times.

    Args:
        char: Character or string to repeat
        n: Number of times to repeat

    Returns:
        Repeated string
    """
    return char * n


def write_html(html_str: str, file_path: str) -> None:
    """
    Write HTML content to a file with proper encoding.

    Args:
        html_str: HTML content to write
        file_path: Path to output file

    Raises:
        IOError: If file cannot be written
    """
    try:
        with open(file_path, mode="w", encoding='utf-8') as html_file:
            html_file.write(html_str)
        # logger.info(f"Successfully wrote {file_path}")
    except IOError as ioe:
        logger.error(f"Error writing file {file_path}: {ioe}")


def load_json(file_path: str) -> dict:
    """
    Load and parse a JSON file.

    Args:
        file_path: Path to JSON file

    Returns:
        Dictionary containing parsed JSON data

    Raises:
        FileNotFoundError: If file doesn't exist
        json.JSONDecodeError: If JSON is invalid
    """
    try:
        with open(file_path, mode='r', encoding='utf-8') as file:
            return json.load(file)
    except FileNotFoundError:
        logger.error(f"Error: File not found: {file_path}")
        raise
    except json.JSONDecodeError as jsone:
        logger.error(f"Error: Invalid JSON in {file_path}: {jsone}")
        raise


def load_description() -> dict[str, dict[str, str]]:
    """
    Load ISO20022 repository description from JSON file.

    Returns:
        Dictionary containing repository description
    """
    logger.info("Loading repository description from JSON file")
    return load_json(REPO_DESC_FILE)


def validate_paths() -> bool:
    """
    Validate that all required paths and files exist.

    Returns:
        True if all paths are valid, False otherwise
    """
    required_paths = {
        'Schema Directory': SCHEMA_DIR,
        'Output Directory': OUTPUT_DIR,
        'Previous Version Directory': PREVIOUS_VERSION_DIR,
        'ISO20022 Model': ISO20022_MODEL,
        'Repository Description File': REPO_DESC_FILE
    }

    all_valid = True
    for name, path in required_paths.items():
        if not os.path.exists(path):
            logger.error(f"Error: {name} not found: {path}")
            all_valid = False

    return all_valid


def setup_output_directory() -> None:
    """
    Ensure output directory exists and is writable.

    Raises:
        OSError: If directory cannot be created or is not writable
    
    """
    try:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        # Test if directory is writable
        test_file = os.path.join(OUTPUT_DIR, '.test')
        try:
            with open(test_file, 'w', encoding='UTF-8') as f:
                f.write('test')
            os.remove(test_file)
        except IOError as ioe:
            logger.error(f"Error: Output directory is not writable: {ioe}")
            raise
    except OSError as ose:
        logger.critical(f"Error creating output directory: {ose}")
        raise

def process_schema(xs: xmlschema.XMLSchema, endpoint_elts: dict[str, xmlschema.XsdElement]):
    # Process main message
    global iso_repository_description


    generate_word_message_table_from_schema(
        SCHEMA_DIR, xs, OUTPUT_DIR)  # type: ignore

    if GENERATE_ENDPOINTS:
        for element_type, endpoint_elt in list(endpoint_elts.items()):
            logger.info("=" * 60)
            logger.info("-" * 16+f"Processing endpoint {element_type}")
            generate_word_message_table_xsd_endpoint(
                xs.name, OUTPUT_DIR, element_type, endpoint_elt)  # type: ignore
            logger.info("-" * 16+f"Processing endpoint {element_type} DONE")

def main() -> None:
    """
    Main execution function with proper error handling.
    """
    try:
        logger.info("=" * 132)
        logger.info(
            "Generate HTML message / components tables from ISO20022 repository")
        logger.info(
            "Needs ISO20022 repository downloadable from ISO20022 website and XML Schema folder to work with")
        logger.info("=" * 132)

        # Validate environment
        if not validate_paths():
            logger.critical(
                "Error: Invalid paths detected. Please check configuration.")
            return
        setup_output_directory()

        # Print configuration
        logger.info(f"SCHEMA_DIR: {SCHEMA_DIR}")
        logger.info(f"OUTPUT_DIR: {OUTPUT_DIR}")
        logger.info(f"ISO20022_MODEL: {ISO20022_MODEL}")

        # Initialize required data
        global endpoints, iso_repository_description, bs_previous_html, element_tags, message_tags
        #Load static information
        endpoints = load_endpoints()
        iso_repository_description = load_description()
        elements_tags, message_tags = parse_repository(ISO20022_MODEL)

        # Process whole directory

        for file in os.listdir(SCHEMA_DIR):
            if file.endswith(".xsd"):
                schema_name = file.split('.xsd')[0]
                logger.info("=" * 80)
                logger.info(f"Processing schema: {schema_name}...")
                # Prepare data from XMLSchema
                xs = xmlschema.XMLSchema(os.path.join(SCHEMA_DIR,file))
                doc = xs.elements['Document']
                endpoints_elts: dict[str,
                                     xmlschema.validators.elements.XsdElement] = {}
                if (doc is None):
                    logger.error(f"Error: Document element not found in schema: {schema_name}... skipping.")
                else:
                    endpoints_elts = collect_all_endpoints_elements(
                        doc, endpoints_elts)
                    process_schema(xs, endpoints_elts)
                    logger.info(f"Processing schema: {schema_name}... DONE")
                    logger.info("=" * 80)
    except Exception as ge:
        logger.critical(f"An error occurred: {ge}")
        traceback.print_exc()

# Load the config file
with open(LOG_CONFIGURATION_FILE, 'rt') as f:
    config = yaml.safe_load(f.read())

# Configure the logging module with the config file
logging.config.dictConfig(config)
logger = logging.getLogger('development')

if __name__ == '__main__':
    main()

